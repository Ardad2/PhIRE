"""
Batch Google Scholar Scraper
Runs multiple queries, saves per-query CSVs, and builds a
deduplicated master papers.docx + master_papers.csv.

Requirements:
    pip install requests beautifulsoup4 python-docx lxml

Usage:
    python scrape_scholar_batch.py
"""

import csv
import os
import random
import re
import sys
import time

import requests
from bs4 import BeautifulSoup

# ── Output settings ───────────────────────────────────────────────────────────
OUTPUT_DIR   = "scholar_results"   # folder for per-query CSVs
MASTER_CSV   = "master_papers.csv"
MASTER_DOCX  = "master_papers.docx"

# ── Global search settings ────────────────────────────────────────────────────
YEAR_LOW  = 2022
NUM_PAGES = 24
PAGE_SIZE = 10
BASE_URL  = "https://scholar.google.com/scholar"

DELAY_MIN = 8    # seconds between page fetches
DELAY_MAX = 20
QUERY_DELAY_MIN = 30   # extra pause between queries (be polite)
QUERY_DELAY_MAX = 60

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/123.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "en-US,en;q=0.9",
}

# ── All queries ───────────────────────────────────────────────────────────────
# Each entry: (label_for_filename, query_string)
QUERIES = [
    # --- Topic-based ---
    ("merge_trees_scalar_field",                '"merge trees" scalar field'),
    ("merge_tree_interpolation",                '"merge tree" interpolation scalar field'),
    ("merge_tree_reconstruction",               '"merge tree" reconstruction scalar field'),
    ("merge_tree_distance",                     '"merge tree" distance scalar field'),
    ("merge_tree_barycenter",                   '"merge tree" barycenter'),
    ("merge_tree_geodesic",                     '"merge tree" geodesic'),
    ("merge_tree_edit_distance",                '"merge tree" edit distance'),
    ("merge_tree_scivis",                       '"merge tree" scientific visualization'),
    ("persistence_diagram_interpolation",       '"persistence diagram" scalar field interpolation'),
    ("persistence_diagram_reconstruction",      '"persistence diagram" scalar field reconstruction'),
    ("topological_descriptors_comparison",      '"topological descriptors" scalar field comparison'),
    ("contour_tree_comparison",                 '"contour tree" scalar field comparison'),
    ("reeb_graph_comparison",                   '"Reeb graph" scalar field comparison'),
    ("morse_smale_comparison",                  '"Morse-Smale" scalar field comparison'),
    ("topology_aware_interpolation",            '"topology aware" scalar field interpolation'),
    ("topology_aware_superresolution",          '"topology aware" scientific super-resolution'),
    ("scivis_merge_persistence",                '"scientific visualization" merge trees persistence diagrams'),
    # --- Author-centered ---
    ("author_tierny",                           'Julien Tierny merge trees persistence diagrams'),
    ("author_bei_wang",                         'Bei Wang scalar field topological descriptors'),
    ("author_hotz",                             'Ingrid Hotz merge trees scalar fields'),
    ("author_masood",                           'Talha Bin Masood merge trees scalar fields'),
    ("author_rasheed",                          'Farhan Rasheed merge trees'),
    ("author_kissi_tierny",                     'Mohamed Kissi Julien Tierny scalar fields'),
]


# ── Scraping helpers ──────────────────────────────────────────────────────────

def fetch_page(session: requests.Session, query: str, start: int) -> BeautifulSoup | None:
    params = {
        "q":      query,
        "hl":     "en",
        "as_sdt": "0,19",
        "as_ylo": str(YEAR_LOW),
        "start":  str(start),
    }
    try:
        resp = session.get(BASE_URL, params=params, headers=HEADERS, timeout=30)
        resp.raise_for_status()
    except requests.RequestException as e:
        print(f"    [ERROR] {e}")
        return None

    soup = BeautifulSoup(resp.text, "html.parser")

    if soup.find("form", {"action": re.compile(r"sorry")}):
        print("\n  *** CAPTCHA detected ***")
        print("  Open https://scholar.google.com in a browser, solve the CAPTCHA,")
        print("  then press Enter here to retry.")
        input("  Press Enter when ready… ")
        return fetch_page(session, query, start)

    return soup


def parse_results(soup: BeautifulSoup) -> list[dict]:
    papers = []
    for result in soup.select(".gs_r.gs_or.gs_scl"):
        title_tag = result.select_one(".gs_rt a")
        if title_tag:
            title = title_tag.get_text(separator=" ", strip=True)
            url   = title_tag.get("href", "")
        else:
            title_tag_nl = result.select_one(".gs_rt")
            title = title_tag_nl.get_text(separator=" ", strip=True) if title_tag_nl else "N/A"
            url   = ""

        meta_tag  = result.select_one(".gs_a")
        meta_text = meta_tag.get_text(" ", strip=True) if meta_tag else ""
        parts     = [p.strip() for p in meta_text.split(" - ")]
        authors   = parts[0] if parts else "N/A"
        years     = re.findall(r"\b(20\d{2}|19\d{2})\b", meta_text)
        year      = years[-1] if years else "N/A"

        papers.append({"title": title, "authors": authors, "year": year, "url": url})
    return papers


def scrape_query(session: requests.Session, query: str, label: str) -> list[dict]:
    all_papers: list[dict] = []
    print(f"\n{'─'*60}")
    print(f"  Query: {query}")
    print(f"  Label: {label}")
    print(f"{'─'*60}")

    for page_num in range(NUM_PAGES):
        start = page_num * PAGE_SIZE
        print(f"  Page {page_num+1:>2}/{NUM_PAGES}  (start={start:>3}) … ", end="", flush=True)

        soup = fetch_page(session, query, start)
        if soup is None:
            print("SKIPPED")
            continue

        results = parse_results(soup)
        all_papers.extend(results)
        print(f"{len(results)} results  (total: {len(all_papers)})")

        # Stop early if Scholar returned no results (end of list)
        if len(results) == 0:
            print("  No more results for this query, moving on.")
            break

        if page_num < NUM_PAGES - 1:
            delay = random.uniform(DELAY_MIN, DELAY_MAX)
            print(f"    ↳ waiting {delay:.1f}s …")
            time.sleep(delay)

    return all_papers


# ── CSV helpers ───────────────────────────────────────────────────────────────

def write_csv(papers: list[dict], path: str) -> None:
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["title", "authors", "year", "url"])
        writer.writeheader()
        writer.writerows(papers)


def read_csv(path: str) -> list[dict]:
    papers = []
    with open(path, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            papers.append(row)
    return papers


# ── Deduplication ─────────────────────────────────────────────────────────────

def deduplicate(papers: list[dict]) -> list[dict]:
    """Deduplicate by URL first, then by normalised title."""
    seen_urls   = set()
    seen_titles = set()
    unique = []
    for p in papers:
        url   = p.get("url", "").strip()
        title = re.sub(r'\W+', '', p.get("title", "").lower())
        if url and url in seen_urls:
            continue
        if title and title in seen_titles:
            continue
        if url:
            seen_urls.add(url)
        if title:
            seen_titles.add(title)
        unique.append(p)
    return unique


# ── DOCX output ───────────────────────────────────────────────────────────────

def write_docx(papers: list[dict], path: str, title_text: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Document title
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(title_text)
    r.bold = True
    r.font.size = Pt(16)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run(f"{len(papers)} papers  |  year ≥ {YEAR_LOW}  |  deduplicated across all queries").font.size = Pt(9)

    doc.add_paragraph()

    for i, paper in enumerate(papers, start=1):
        # Title line
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        tr = p.add_run(paper["title"])
        tr.bold = True
        tr.font.size = Pt(11)

        # Authors | Year
        meta = doc.add_paragraph()
        meta.paragraph_format.left_indent = Pt(20)
        mr = meta.add_run(f"{paper['authors']}  |  {paper['year']}")
        mr.font.size = Pt(10)
        mr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # URL as hyperlink
        url = paper.get("url", "").strip()
        url_para = doc.add_paragraph()
        url_para.paragraph_format.left_indent = Pt(20)
        if url:
            rId = url_para.part.relate_to(
                url,
                "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                is_external=True,
            )
            hyperlink = OxmlElement("w:hyperlink")
            hyperlink.set(qn("r:id"), rId)
            wr = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), "Hyperlink")
            rPr.append(rStyle)
            wr.append(rPr)
            t = OxmlElement("w:t")
            t.text = url
            wr.append(t)
            hyperlink.append(wr)
            url_para._p.append(hyperlink)
        else:
            url_para.add_run("URL: N/A").font.size = Pt(10)

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    doc.save(path)
    print(f"  DOCX saved → {path}  ({len(papers)} papers)")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    session      = requests.Session()
    all_papers: list[dict] = []

    total_queries = len(QUERIES)
    for qi, (label, query) in enumerate(QUERIES, start=1):
        print(f"\n{'='*60}")
        print(f"  QUERY {qi}/{total_queries}: {label}")
        print(f"{'='*60}")

        csv_path = os.path.join(OUTPUT_DIR, f"{label}.csv")

        # Resume: skip if already scraped
        if os.path.exists(csv_path):
            print(f"  [SKIP] Already scraped → loading {csv_path}")
            papers = read_csv(csv_path)
        else:
            papers = scrape_query(session, query, label)
            write_csv(papers, csv_path)
            print(f"  Per-query CSV saved → {csv_path}  ({len(papers)} papers)")

        all_papers.extend(papers)

        # Polite pause between queries (skip after last one)
        if qi < total_queries:
            delay = random.uniform(QUERY_DELAY_MIN, QUERY_DELAY_MAX)
            print(f"\n  ⏸  Pausing {delay:.0f}s before next query …")
            time.sleep(delay)

    # Deduplicate and write master outputs
    print(f"\n{'='*60}")
    print(f"  Deduplicating {len(all_papers)} total records …")
    unique = deduplicate(all_papers)
    print(f"  {len(unique)} unique papers after deduplication")

    write_csv(unique, MASTER_CSV)
    print(f"  Master CSV saved → {MASTER_CSV}")

    write_docx(
        unique,
        MASTER_DOCX,
        'Scholar Papers: Merge Trees / Topological Methods (2022–)',
    )

    print("\nAll done ✓")
    print(f"  Per-query CSVs  → ./{OUTPUT_DIR}/")
    print(f"  Master CSV      → {MASTER_CSV}")
    print(f"  Master DOCX     → {MASTER_DOCX}")


if __name__ == "__main__":
    main()
