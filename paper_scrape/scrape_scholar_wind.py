"""
Batch Google Scholar Scraper — Wind SR / Scientific SR / Metrics / PhIRE
Uses the `scholarly` library for robust Scholar access with built-in rate limiting.

Requirements:
    pip install scholarly python-docx lxml

Usage:
    python3 scrape_scholar_wind.py
"""

import csv
import os
import re
import time
from collections import Counter

from scholarly import scholarly, ProxyGenerator

# ── Output settings ───────────────────────────────────────────────────────────
OUTPUT_DIR   = "scholar_results_wind"
MASTER_CSV   = "master_papers_wind.csv"
MASTER_DOCX  = "master_papers_wind.docx"

# ── Global search settings ────────────────────────────────────────────────────
YEAR_LOW   = 2018
MAX_RESULTS = 240   # 24 pages × 10 results

# ── Queries by bucket ─────────────────────────────────────────────────────────
QUERIES = [

    # ── Bucket A: Wind-field / atmospheric SR / downscaling ───────────────────
    ("wind_field_sr",                   '"wind field" super-resolution',                            "A"),
    ("wind_field_sr_hyphen",            '"wind-field" super-resolution',                            "A"),
    ("wind_speed_sr",                   '"wind speed" super-resolution',                            "A"),
    ("wind_velocity_sr",                '"wind velocity" super-resolution',                         "A"),
    ("atmospheric_sr_wind",             '"atmospheric" super-resolution wind',                      "A"),
    ("meteorological_field_sr",         '"meteorological field" super-resolution',                  "A"),
    ("climate_downscaling_dl_wind",     '"climate downscaling" deep learning wind',                 "A"),
    ("wind_downscaling_dl",             '"wind downscaling" deep learning',                         "A"),
    ("spatial_downscaling_wind_nn",     '"spatial downscaling" wind field neural network',          "A"),
    ("wrf_wind_sr",                     '"WRF" wind super-resolution',                             "A"),
    ("wind_field_recon_dl",             '"wind field" reconstruction deep learning',                "A"),
    ("scientific_sr_wind",              '"scientific super-resolution" wind',                       "A"),

    # ── Bucket B: Scalar / scientific-field SR ────────────────────────────────
    ("scalar_field_sr",                 '"scalar field" super-resolution',                          "B"),
    ("scientific_data_sr",              '"scientific data" super-resolution scalar field',          "B"),
    ("scivis_sr_scalar",                '"scientific visualization" super-resolution scalar field', "B"),
    ("physics_informed_sr",             '"physics-informed" scalar field super-resolution',         "B"),
    ("simulation_sr_scalar",            '"simulation" super-resolution scalar field',               "B"),
    ("fluid_field_sr",                  '"fluid field" super-resolution',                           "B"),
    ("flow_field_sr",                   '"flow field" super-resolution',                            "B"),
    ("velocity_field_sr",               '"velocity field" super-resolution',                        "B"),
    ("temperature_field_sr",            '"temperature field" super-resolution',                     "B"),
    ("climate_field_sr",                '"climate field" super-resolution',                         "B"),

    # ── Bucket C: SR metrics and their limitations ────────────────────────────
    ("sr_psnr_ssim_limits",             '"super-resolution" PSNR SSIM limitations',                "C"),
    ("scientific_sr_psnr_ssim",         '"scientific super-resolution" PSNR SSIM',                 "C"),
    ("perceptual_metrics_sr",           '"perceptual metrics" super-resolution scientific data',   "C"),
    ("eval_metrics_sr",                 '"evaluation metrics" super-resolution scientific data',   "C"),
    ("psnr_ssim_scivis",                '"PSNR" SSIM scientific visualization reconstruction',     "C"),
    ("pixelwise_limits",                '"super-resolution" "pixel-wise metrics" limitations',     "C"),
    ("perceptual_distortion_tradeoff",  '"image super-resolution" perceptual distortion tradeoff', "C"),
    ("srgan_perceptual_loss",           '"SRGAN" perceptual loss PSNR',                            "C"),
    ("sr_eval_review",                  '"super-resolution" "evaluation metric" review',           "C"),

    # ── Bucket D: PhIRE / wind-energy / domain justification ─────────────────
    ("phire_wind_sr",                   '"PhIRE" wind super-resolution',                           "D"),
    ("stengel_phire",                   '"Stengel" PhIRE wind',                                    "D"),
    ("wind_energy_downscaling_dl",      '"wind energy" downscaling deep learning',                 "D"),
    ("wind_resource_sr",                '"wind resource" super-resolution',                        "D"),
    ("wind_energy_forecast_downscaling",'"wind energy forecasting" spatial downscaling',           "D"),
    ("atm_downscaling_wind_energy",     '"atmospheric downscaling" wind energy',                   "D"),
    ("hires_wind_dl",                   '"high-resolution wind field" deep learning',              "D"),
]

BUCKET_NAMES = {
    "A": "Bucket A — Wind-field / atmospheric SR / downscaling",
    "B": "Bucket B — Scalar / scientific-field SR",
    "C": "Bucket C — SR metrics and their limitations",
    "D": "Bucket D — PhIRE / wind-energy / domain justification",
}


# ── Scraping ──────────────────────────────────────────────────────────────────

def scrape_query(query: str, label: str, max_results: int = MAX_RESULTS) -> list[dict]:
    print(f"\n{'─'*60}")
    print(f"  Query : {query}")
    print(f"  Label : {label}")
    print(f"{'─'*60}")

    papers = []
    try:
        search_gen = scholarly.search_pubs(query)
        for i, result in enumerate(search_gen):
            if i >= max_results:
                break

            bib  = result.get("bib", {})
            title   = bib.get("title", "N/A")
            authors = ", ".join(bib.get("author", [])) or "N/A"
            year    = str(bib.get("pub_year", "N/A"))
            url     = result.get("pub_url", "") or result.get("eprint_url", "") or ""

            # Filter by year
            if year.isdigit() and int(year) < YEAR_LOW:
                continue

            papers.append({
                "title":   title,
                "authors": authors,
                "year":    year,
                "url":     url,
            })

            if (i + 1) % 10 == 0:
                print(f"  Fetched {i+1} results so far…", flush=True)

    except Exception as e:
        print(f"  [ERROR] {e}")

    print(f"  Done — {len(papers)} papers collected.")
    return papers


# ── CSV helpers ───────────────────────────────────────────────────────────────

def write_csv(papers: list[dict], path: str, extra_fields: list[str] | None = None) -> None:
    fields = ["title", "authors", "year", "url"] + (extra_fields or [])
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(papers)


def read_csv(path: str) -> list[dict]:
    with open(path, newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


# ── Deduplication ─────────────────────────────────────────────────────────────

def deduplicate(papers: list[dict]) -> list[dict]:
    seen_urls   = set()
    seen_titles = set()
    unique = []
    for p in papers:
        url   = p.get("url", "").strip()
        title = re.sub(r'\W+', '', p.get("title", "").lower())
        if url and url in seen_urls:
            continue
        if title and title in seen_titles:
            continue
        if url:
            seen_urls.add(url)
        if title:
            seen_titles.add(title)
        unique.append(p)
    return unique


# ── DOCX output ───────────────────────────────────────────────────────────────

def write_docx(papers: list[dict], path: str, doc_title: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tp.add_run(doc_title)
    r.bold = True
    r.font.size = Pt(16)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.add_run(
        f"{len(papers)} unique papers  |  year ≥ {YEAR_LOW}  |  "
        f"deduplicated across {len(QUERIES)} queries / 4 buckets"
    ).font.size = Pt(9)
    doc.add_paragraph()

    bucket_order = ["A", "B", "C", "D"]
    by_bucket: dict[str, list[dict]] = {b: [] for b in bucket_order}
    for p in papers:
        by_bucket.setdefault(p.get("bucket", "?"), []).append(p)

    global_idx = 1
    for bucket_key in bucket_order:
        bucket_papers = by_bucket.get(bucket_key, [])
        if not bucket_papers:
            continue

        bh = doc.add_paragraph()
        bhr = bh.add_run(BUCKET_NAMES[bucket_key])
        bhr.bold = True
        bhr.font.size = Pt(13)
        bh.paragraph_format.space_before = Pt(14)
        bh.paragraph_format.space_after  = Pt(4)

        for paper in bucket_papers:
            p = doc.add_paragraph()
            p.add_run(f"{global_idx}. ").bold = True
            tr = p.add_run(paper["title"])
            tr.bold = True
            tr.font.size = Pt(11)

            meta = doc.add_paragraph()
            meta.paragraph_format.left_indent = Pt(20)
            mr = meta.add_run(f"{paper['authors']}  |  {paper['year']}")
            mr.font.size = Pt(10)
            mr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            url = paper.get("url", "").strip()
            url_para = doc.add_paragraph()
            url_para.paragraph_format.left_indent = Pt(20)
            if url:
                rId = url_para.part.relate_to(
                    url,
                    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
                    is_external=True,
                )
                hyperlink = OxmlElement("w:hyperlink")
                hyperlink.set(qn("r:id"), rId)
                wr = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rStyle = OxmlElement("w:rStyle")
                rStyle.set(qn("w:val"), "Hyperlink")
                rPr.append(rStyle)
                wr.append(rPr)
                t = OxmlElement("w:t")
                t.text = url
                wr.append(t)
                hyperlink.append(wr)
                url_para._p.append(hyperlink)
            else:
                url_para.add_run("URL: N/A").font.size = Pt(10)

            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            global_idx += 1

    doc.save(path)
    print(f"  DOCX saved → {path}  ({len(papers)} papers)")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    all_papers: list[dict] = []

    total = len(QUERIES)
    for qi, (label, query, bucket) in enumerate(QUERIES, start=1):
        print(f"\n{'='*60}")
        print(f"  QUERY {qi}/{total}  [{BUCKET_NAMES[bucket]}]")
        print(f"{'='*60}")

        csv_path = os.path.join(OUTPUT_DIR, f"{label}.csv")

        if os.path.exists(csv_path):
            print(f"  [SKIP] Already scraped → loading {csv_path}")
            papers = read_csv(csv_path)
            for p in papers:
                p.setdefault("bucket", bucket)
            all_papers.extend(papers)
            continue

        papers = scrape_query(query, label)
        for p in papers:
            p["bucket"] = bucket

        if papers:
            write_csv(papers, csv_path, extra_fields=["bucket"])
            print(f"  Saved → {csv_path}  ({len(papers)} papers)")
        else:
            print(f"  [WARN] No papers — will retry next run.")

        all_papers.extend(papers)

    # Deduplicate & write outputs
    print(f"\n{'='*60}")
    print(f"  Deduplicating {len(all_papers)} total records …")
    unique = deduplicate(all_papers)
    print(f"  {len(unique)} unique papers after deduplication")

    write_csv(unique, MASTER_CSV, extra_fields=["bucket"])
    print(f"  Master CSV → {MASTER_CSV}")

    write_docx(unique, MASTER_DOCX, "Scholar Papers: Wind SR / Scientific SR / Metrics / PhIRE")

    print(f"\n{'─'*60}")
    print("  Summary by bucket:")
    counts = Counter(p.get("bucket", "?") for p in unique)
    for b in ["A", "B", "C", "D"]:
        print(f"    {BUCKET_NAMES[b]}: {counts.get(b, 0)} papers")
    print(f"{'─'*60}")

    print("\nAll done ✓")


if __name__ == "__main__":
    main()